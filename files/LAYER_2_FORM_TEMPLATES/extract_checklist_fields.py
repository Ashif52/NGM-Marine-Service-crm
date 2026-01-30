"""
Extract all fields from forms in the Checklists subfolder.
Supports both Word documents (.doc/.docx) and Excel spreadsheets (.xls/.xlsx)
"""

import os
import json
from pathlib import Path
from typing import Dict, List, Any
import re

try:
    from docx import Document
    from docx.table import Table
    from docx.text.paragraph import Paragraph
except ImportError:
    print("Installing python-docx...")
    os.system("pip install python-docx")
    from docx import Document
    from docx.table import Table
    from docx.text.paragraph import Paragraph

try:
    import openpyxl
except ImportError:
    print("Installing openpyxl...")
    os.system("pip install openpyxl")
    import openpyxl

try:
    import pandas as pd
except ImportError:
    print("Installing pandas...")
    os.system("pip install pandas")
    import pandas as pd


def extract_fields_from_docx(file_path: str) -> Dict[str, Any]:
    """Extract fields from a .docx file"""
    try:
        doc = Document(file_path)
        fields = {
            "filename": os.path.basename(file_path),
            "type": "docx",
            "paragraphs": [],
            "tables": [],
            "form_fields": [],
            "headers": [],
            "potential_fields": []
        }
        
        # Extract paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                fields["paragraphs"].append(text)
                
                # Identify potential form fields (lines ending with : or containing ___)
                if ':' in text or '___' in text or '[]' in text or '☐' in text:
                    fields["potential_fields"].append(text)
        
        # Extract tables
        for table_idx, table in enumerate(doc.tables):
            table_data = {
                "table_number": table_idx + 1,
                "rows": len(table.rows),
                "columns": len(table.columns),
                "headers": [],
                "data": []
            }
            
            for row_idx, row in enumerate(table.rows):
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    row_data.append(cell_text)
                    
                if row_idx == 0:
                    table_data["headers"] = row_data
                else:
                    table_data["data"].append(row_data)
            
            fields["tables"].append(table_data)
        
        return fields
    except Exception as e:
        return {
            "filename": os.path.basename(file_path),
            "type": "docx",
            "error": str(e)
        }


def extract_fields_from_xlsx(file_path: str) -> Dict[str, Any]:
    """Extract fields from an Excel file"""
    try:
        workbook = openpyxl.load_workbook(file_path, data_only=True)
        fields = {
            "filename": os.path.basename(file_path),
            "type": "xlsx",
            "sheets": []
        }
        
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            sheet_data = {
                "sheet_name": sheet_name,
                "headers": [],
                "columns": [],
                "rows": sheet.max_row,
                "cols": sheet.max_column,
                "data_sample": []
            }
            
            # Extract first row as potential headers
            first_row = []
            for col in range(1, min(sheet.max_column + 1, 50)):  # Limit to 50 columns
                cell_value = sheet.cell(row=1, column=col).value
                if cell_value:
                    first_row.append(str(cell_value).strip())
                else:
                    first_row.append("")
            
            sheet_data["headers"] = first_row
            
            # Extract sample data (first 10 rows)
            for row in range(1, min(sheet.max_row + 1, 11)):
                row_data = []
                for col in range(1, min(sheet.max_column + 1, 50)):
                    cell_value = sheet.cell(row=row, column=col).value
                    if cell_value:
                        row_data.append(str(cell_value).strip())
                    else:
                        row_data.append("")
                sheet_data["data_sample"].append(row_data)
            
            fields["sheets"].append(sheet_data)
        
        return fields
    except Exception as e:
        return {
            "filename": os.path.basename(file_path),
            "type": "xlsx",
            "error": str(e)
        }


def extract_fields_from_xls(file_path: str) -> Dict[str, Any]:
    """Extract fields from an old Excel file (.xls) using pandas"""
    try:
        # Read all sheets
        excel_file = pd.ExcelFile(file_path)
        fields = {
            "filename": os.path.basename(file_path),
            "type": "xls",
            "sheets": []
        }
        
        for sheet_name in excel_file.sheet_names:
            df = pd.read_excel(file_path, sheet_name=sheet_name, header=None)
            
            sheet_data = {
                "sheet_name": sheet_name,
                "headers": [],
                "rows": len(df),
                "cols": len(df.columns),
                "data_sample": []
            }
            
            # Get first row as potential headers
            if len(df) > 0:
                first_row = df.iloc[0].fillna('').astype(str).tolist()
                sheet_data["headers"] = first_row
            
            # Get sample data (first 10 rows)
            for idx in range(min(len(df), 10)):
                row_data = df.iloc[idx].fillna('').astype(str).tolist()
                sheet_data["data_sample"].append(row_data)
            
            fields["sheets"].append(sheet_data)
        
        return fields
    except Exception as e:
        return {
            "filename": os.path.basename(file_path),
            "type": "xls",
            "error": str(e)
        }


def process_checklists_folder():
    """Process all forms in the Checklists folder"""
    base_path = Path(__file__).parent / "Checklists"
    
    if not base_path.exists():
        print(f"Error: Checklists folder not found at {base_path}")
        return
    
    all_forms = []
    
    # Get all files
    files = sorted(base_path.glob("*"))
    
    print(f"Processing {len(files)} files from Checklists folder...\n")
    
    for file_path in files:
        if file_path.is_file():
            print(f"Processing: {file_path.name}")
            
            ext = file_path.suffix.lower()
            
            if ext == '.docx':
                form_data = extract_fields_from_docx(str(file_path))
            elif ext == '.doc':
                # Try to handle .doc files - might need win32com on Windows
                print(f"  Warning: .doc format might not be fully supported. Consider converting to .docx")
                # Skip .doc files for now or try with win32com
                form_data = {
                    "filename": file_path.name,
                    "type": "doc",
                    "note": "Old Word format - please convert to .docx for full extraction"
                }
            elif ext in ['.xls', '.xlsx']:
                if ext == '.xlsx':
                    form_data = extract_fields_from_xlsx(str(file_path))
                else:
                    form_data = extract_fields_from_xls(str(file_path))
            else:
                continue
            
            all_forms.append(form_data)
    
    # Save results to JSON
    output_file = Path(__file__).parent / "checklists_fields_extraction.json"
    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump(all_forms, f, indent=2, ensure_ascii=False)
    
    print(f"\n✅ Extraction complete! Results saved to: {output_file}")
    
    # Generate summary
    generate_summary(all_forms)
    
    return all_forms


def generate_summary(all_forms: List[Dict[str, Any]]):
    """Generate a readable summary of extracted fields"""
    output_file = Path(__file__).parent / "checklists_fields_summary.md"
    
    with open(output_file, 'w', encoding='utf-8') as f:
        f.write("# Checklists Folder - Form Fields Extraction Summary\n\n")
        f.write(f"**Total Forms Analyzed**: {len(all_forms)}\n\n")
        f.write("---\n\n")
        
        for idx, form in enumerate(all_forms, 1):
            f.write(f"## {idx}. {form['filename']}\n\n")
            f.write(f"**Type**: {form['type'].upper()}\n\n")
            
            if 'error' in form:
                f.write(f"❌ **Error**: {form['error']}\n\n")
                continue
            
            if form['type'] in ['docx']:
                # Document analysis
                f.write(f"### Document Structure\n\n")
                f.write(f"- **Paragraphs**: {len(form.get('paragraphs', []))}\n")
                f.write(f"- **Tables**: {len(form.get('tables', []))}\n")
                f.write(f"- **Potential Form Fields**: {len(form.get('potential_fields', []))}\n\n")
                
                # List potential form fields
                if form.get('potential_fields'):
                    f.write(f"### Identified Form Fields\n\n")
                    for field in form['potential_fields'][:20]:  # Limit to first 20
                        f.write(f"- {field}\n")
                    
                    if len(form['potential_fields']) > 20:
                        f.write(f"\n*...and {len(form['potential_fields']) - 20} more fields*\n")
                    f.write("\n")
                
                # Table information
                if form.get('tables'):
                    f.write(f"### Tables\n\n")
                    for table in form['tables']:
                        f.write(f"**Table {table['table_number']}**: {table['rows']} rows × {table['columns']} columns\n\n")
                        if table['headers']:
                            f.write(f"Headers: {' | '.join(table['headers'][:10])}\n\n")
            
            elif form['type'] in ['xls', 'xlsx']:
                # Spreadsheet analysis
                f.write(f"### Spreadsheet Structure\n\n")
                f.write(f"- **Sheets**: {len(form.get('sheets', []))}\n\n")
                
                for sheet in form.get('sheets', []):
                    f.write(f"#### Sheet: {sheet['sheet_name']}\n\n")
                    f.write(f"- **Dimensions**: {sheet['rows']} rows × {sheet['cols']} columns\n")
                    
                    if sheet.get('headers'):
                        non_empty_headers = [h for h in sheet['headers'] if h]
                        if non_empty_headers:
                            f.write(f"- **Column Headers**:\n")
                            for header in non_empty_headers[:15]:
                                f.write(f"  - {header}\n")
                            if len(non_empty_headers) > 15:
                                f.write(f"  - *...and {len(non_empty_headers) - 15} more columns*\n")
                    f.write("\n")
            
            elif form['type'] == 'doc':
                f.write(f"⚠️ {form.get('note', 'Old format')}\n\n")
            
            f.write("---\n\n")
    
    print(f"📄 Summary saved to: {output_file}")


if __name__ == "__main__":
    process_checklists_folder()
